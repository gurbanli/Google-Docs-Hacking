import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document


def createDocument(text):
    document = Document()
    document.add_heading('Database Questions', 0)
    document.add_paragraph(text)
    document.save('Database.docx')


changeScroll = """
    followers = document.querySelector('div.kix-appview-editor');
    followers.scrollTo(0, {});
    """
chrome_options = Options()
chrome_options.add_argument('--headless')
chrome_options.add_argument('--start-maximized')
driver = webdriver.Chrome(options=chrome_options)
scroll = 173779
count = 1
driver.get("https://docs.google.com/document/d/1fPQ8IXt2XAzZv-NzW7rzyhviFVvCu5KJcBp09dL4Et8/preview")
time.sleep(10)
scrolled = changeScroll.format(scroll)
driver.execute_script(scrolled)
time.sleep(2)
driver.set_window_size(1920, scroll)
time.sleep(2)
element = driver.find_element_by_tag_name('body')
createDocument(element.text)
driver.quit()
